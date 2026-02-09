"""
MCP 文件解析服务 - 提供多种文件格式的解析能力
"""
import io
import json
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, List, Dict

# PDF 解析器
def parse_pdf_pymupdf(file_bytes: bytes) -> dict:
    """PyMuPDF - 提取所有文字层（包括隐藏）"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            pages.append({
                "page": page_num + 1,
                "text": text,
                "char_count": len(text)
            })
        doc.close()
        return {
            "parser": "pymupdf",
            "success": True,
            "total_pages": len(pages),
            "pages": pages,
            "extracts_hidden": True
        }
    except Exception as e:
        return {"parser": "pymupdf", "success": False, "error": str(e)}


def parse_pdf_pdfplumber(file_bytes: bytes) -> dict:
    """pdfplumber - 结构化文本提取"""
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                tables = page.extract_tables()
                pages.append({
                    "page": page_num + 1,
                    "text": text,
                    "tables": tables,
                    "char_count": len(text)
                })
        return {
            "parser": "pdfplumber",
            "success": True,
            "total_pages": len(pages),
            "pages": pages,
            "extracts_hidden": True
        }
    except Exception as e:
        return {"parser": "pdfplumber", "success": False, "error": str(e)}


def parse_pdf_ocr(file_bytes: bytes) -> dict:
    """pdf2image + OCR - 仅识别可见内容"""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract

        images = convert_from_bytes(file_bytes)
        pages = []
        for page_num, image in enumerate(images):
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            pages.append({
                "page": page_num + 1,
                "text": text,
                "char_count": len(text)
            })
        return {
            "parser": "pdf2image_ocr",
            "success": True,
            "total_pages": len(pages),
            "pages": pages,
            "extracts_hidden": False,
            "note": "OCR 只能识别可见内容，隐藏文本不会被提取"
        }
    except Exception as e:
        return {"parser": "pdf2image_ocr", "success": False, "error": str(e)}


# DOCX 解析器
def parse_docx_python_docx(file_bytes: bytes) -> dict:
    """python-docx - 包含隐藏文本"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))

        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else None
            })

        # 提取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return {
            "parser": "python-docx",
            "success": True,
            "paragraphs": paragraphs,
            "tables": tables,
            "extracts_hidden": True
        }
    except Exception as e:
        return {"parser": "python-docx", "success": False, "error": str(e)}


def parse_docx_mammoth(file_bytes: bytes) -> dict:
    """mammoth - 转换为 HTML/文本"""
    try:
        import mammoth
        result = mammoth.convert_to_html(io.BytesIO(file_bytes))
        text_result = mammoth.extract_raw_text(io.BytesIO(file_bytes))

        return {
            "parser": "mammoth",
            "success": True,
            "html": result.value,
            "text": text_result.value,
            "messages": [str(m) for m in result.messages],
            "extracts_hidden": False
        }
    except Exception as e:
        return {"parser": "mammoth", "success": False, "error": str(e)}


# XLSX 解析器
def parse_xlsx_openpyxl(file_bytes: bytes, include_hidden: bool = False) -> dict:
    """openpyxl - Excel 文件解析"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes))

        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # 检查工作表是否隐藏
            is_hidden = ws.sheet_state != 'visible'
            if is_hidden and not include_hidden:
                continue

            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(list(row))

            sheets.append({
                "name": sheet_name,
                "hidden": is_hidden,
                "rows": rows
            })

        wb.close()
        return {
            "parser": "openpyxl" + ("_hidden" if include_hidden else ""),
            "success": True,
            "sheets": sheets,
            "extracts_hidden": include_hidden
        }
    except Exception as e:
        return {"parser": "openpyxl", "success": False, "error": str(e)}


# 图片解析器
def parse_image_exiftool(file_bytes: bytes, filename: str = "image") -> dict:
    """exiftool - 提取 EXIF/元数据"""
    try:
        with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as f:
            f.write(file_bytes)
            temp_path = f.name

        result = subprocess.run(
            ["exiftool", "-j", temp_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        Path(temp_path).unlink()

        if result.returncode == 0:
            metadata = json.loads(result.stdout)
            return {
                "parser": "exiftool",
                "success": True,
                "metadata": metadata[0] if metadata else {},
                "extracts_hidden": True
            }
        else:
            return {"parser": "exiftool", "success": False, "error": result.stderr}
    except Exception as e:
        return {"parser": "exiftool", "success": False, "error": str(e)}


def parse_image_ocr(file_bytes: bytes) -> dict:
    """pytesseract - OCR 文字识别"""
    try:
        from PIL import Image
        import pytesseract

        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image, lang='chi_sim+eng')

        return {
            "parser": "pytesseract",
            "success": True,
            "text": text,
            "image_size": image.size,
            "extracts_hidden": False
        }
    except Exception as e:
        return {"parser": "pytesseract", "success": False, "error": str(e)}


def parse_image_pillow_meta(file_bytes: bytes) -> dict:
    """Pillow - 读取图片元数据和注释"""
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS

        image = Image.open(io.BytesIO(file_bytes))

        # 基本信息
        info = {
            "format": image.format,
            "mode": image.mode,
            "size": image.size,
        }

        # EXIF 数据
        exif_data = {}
        if hasattr(image, '_getexif') and image._getexif():
            for tag_id, value in image._getexif().items():
                tag = TAGS.get(tag_id, tag_id)
                try:
                    exif_data[tag] = str(value) if not isinstance(value, (str, int, float)) else value
                except Exception:
                    pass

        # 图片注释/描述
        comments = {}
        if image.info:
            for key, value in image.info.items():
                try:
                    comments[key] = str(value) if not isinstance(value, (str, int, float, bytes)) else value
                except Exception:
                    pass

        return {
            "parser": "pillow_meta",
            "success": True,
            "info": info,
            "exif": exif_data,
            "comments": comments,
            "extracts_hidden": True
        }
    except Exception as e:
        return {"parser": "pillow_meta", "success": False, "error": str(e)}


# 解析器注册表
PARSERS = {
    "pdf": {
        "pymupdf": parse_pdf_pymupdf,
        "pdfplumber": parse_pdf_pdfplumber,
        "pdf2image_ocr": parse_pdf_ocr,
    },
    "docx": {
        "python-docx": parse_docx_python_docx,
        "mammoth": parse_docx_mammoth,
    },
    "xlsx": {
        "openpyxl": lambda b: parse_xlsx_openpyxl(b, include_hidden=False),
        "openpyxl_hidden": lambda b: parse_xlsx_openpyxl(b, include_hidden=True),
    },
    "image": {
        "exiftool": parse_image_exiftool,
        "pytesseract": parse_image_ocr,
        "pillow_meta": parse_image_pillow_meta,
    }
}


def get_file_type(filename: str) -> Optional[str]:
    """根据文件扩展名确定文件类型"""
    ext = Path(filename).suffix.lower()
    type_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".xlsx": "xlsx",
        ".xls": "xlsx",
        ".jpg": "image",
        ".jpeg": "image",
        ".png": "image",
        ".gif": "image",
        ".bmp": "image",
        ".webp": "image",
    }
    return type_map.get(ext)


def parse_file(file_bytes: bytes, filename: str, parser_ids: List[str]) -> List[Dict]:
    """
    使用指定的解析器解析文件

    Args:
        file_bytes: 文件内容
        filename: 文件名（用于判断类型）
        parser_ids: 要使用的解析器 ID 列表

    Returns:
        解析结果列表
    """
    file_type = get_file_type(filename)
    if not file_type:
        return [{"error": f"不支持的文件类型: {filename}"}]

    available_parsers = PARSERS.get(file_type, {})
    results = []

    for parser_id in parser_ids:
        if parser_id in available_parsers:
            parser_func = available_parsers[parser_id]
            # 对于 exiftool，需要传入文件名
            if parser_id == "exiftool":
                result = parser_func(file_bytes, filename)
            else:
                result = parser_func(file_bytes)
            results.append(result)

    return results
